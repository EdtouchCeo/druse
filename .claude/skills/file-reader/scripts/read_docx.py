"""read_docx.py — .docx 파일에서 텍스트 추출"""
import sys
import os

def read_docx(file_path: str) -> str:
    try:
        import docx
    except ImportError:
        print("오류: python-docx가 설치되지 않았습니다. pip install python-docx", file=sys.stderr)
        sys.exit(1)

    if not os.path.exists(file_path):
        print(f"오류: 파일을 찾을 수 없습니다 — {file_path}", file=sys.stderr)
        sys.exit(1)

    doc = docx.Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 표 내용도 추출
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    result = "\n".join(paragraphs)
    if not result.strip():
        print(f"오류: 텍스트를 추출할 수 없습니다 — {file_path}", file=sys.stderr)
        sys.exit(1)

    return result


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("사용법: python read_docx.py <파일경로>", file=sys.stderr)
        sys.exit(1)

    text = read_docx(sys.argv[1])
    print(text)
